"""Dump .docx paragraphs and tables to UTF-8 text file."""
import sys
from pathlib import Path
from docx import Document


def dump(docx_path: str, out_path: str) -> None:
    d = Document(docx_path)
    lines: list[str] = []
    lines.append(f"=== FILE: {docx_path} ===\n")
    lines.append(f"=== PARAGRAPHS ({len(d.paragraphs)}) ===")
    for i, p in enumerate(d.paragraphs):
        if p.text.strip():
            lines.append(f"[P{i}] {p.text}")

    lines.append(f"\n=== TABLES ({len(d.tables)}) ===")
    for ti, t in enumerate(d.tables):
        lines.append(f"\n--- Table #{ti} ({len(t.rows)} rows x {len(t.columns)} cols) ---")
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = "\n".join(p.text for p in cell.paragraphs).strip()
                if txt:
                    lines.append(f"  [T{ti}.R{ri}.C{ci}] {txt}")

    Path(out_path).write_text("\n".join(lines), encoding="utf-8")
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    dump(sys.argv[1], sys.argv[2])
