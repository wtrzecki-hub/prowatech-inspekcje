"""Generate updated PIIB protocol templates (.docx).

Outputs to wzory_PIIB/:
  - Protokol_Kontroli_Rocznej_EW_PIIB_R.docx     (rozszerzona, default)
  - Protokol_Kontroli_Rocznej_EW_PIIB_U.docx     (uproszczona, bez wjazdu)
  - Protokol_Kontroli_5-letniej_EW_PIIB.docx     (jedyna wersja, zawsze z wjazdem)

Plus updates Raport_zmian_wzory_PIIB.docx (dodaje sekcję 9).
"""
from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import _Cell


# Style colors matching the app's protocol-tokens (HEX.graphite800, HEX.graphite200)
# Header fill: graphite800 = 1B2230 (dark navy)
# Border color: graphite200 = DDE3EA (light grey)
TABLE_HEADER_FILL = "1B2230"
TABLE_BORDER_COLOR = "DDE3EA"
TABLE_HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF)
TABLE_ZEBRA_FILL = "F5F7F9"  # very light grey for alternating body rows

# Font face — aplikacja używa 'Arial' (HEX.graphite900 dla tekstu domyślnie)
ARIAL_FONT = "Arial"
GRAPHITE_900_HEX = "0F1520"

# Tabele które MAJĄ nagłówek u góry (header row 0) — kandydaci do stylowania Żeńsko.
# Tabele które są typu "label : value" (np. metryczka, firma serwisowa) — bez header'a, tylko borders.
# Heuristyka: jeśli row 0 ma KRÓTKIE labele jak "Lp.", "Symbol", "Klasa" etc. — header.
HEADER_TABLE_HINTS = {
    "lp.", "symbol", "klasa", "stopień", "wymaganie podstawowe",
    "branża", "rodzaj pomiaru / sprawdzenia", "element, urządzenie",
    "element obiektu", "zalecenia z poprzedniej kontroli", "model",
    "imię i nazwisko", "rodzaj konstrukcji", "zakres czynności",
    "dobry",  # Kryteria oceny PIIB
    "rodzaj pomiaru", "nr protokołu z pomiaru",
}


SOURCE_DIR = Path(r"C:\prowatech-inspekcje\wzory_PIIB")
# Original rocznikowy was deleted — use _R as source (it has unchanged cells E1, E2, E4, E5, E8-15
# with original styling, which we use as deep-copy templates for re-styling modified cells).
ROCZNY_SRC = SOURCE_DIR / "Protokol_Kontroli_Rocznej_EW_PIIB_R.docx"
PIECIOLETNI_SRC = SOURCE_DIR / "Protokol_Kontroli_5-letniej_EW_PIIB.docx"
RAPORT_SRC = SOURCE_DIR / "Raport_zmian_wzory_PIIB.docx"

ROCZNY_R_OUT = SOURCE_DIR / "Protokol_Kontroli_Rocznej_EW_PIIB_R.docx"
ROCZNY_U_OUT = SOURCE_DIR / "Protokol_Kontroli_Rocznej_EW_PIIB_U.docx"
PIECIOLETNI_OUT = SOURCE_DIR / "Protokol_Kontroli_5-letniej_EW_PIIB.docx"  # in place
PIECIOLETNI_O_OUT = SOURCE_DIR / "Protokol_Kontroli_5-letniej_EW_PIIB_O.docx"  # odrębny

W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ---------- helpers ---------------------------------------------------------


def clear_cell(cell: _Cell) -> None:
    """Remove all paragraphs from the cell."""
    tc = cell._tc
    for child in list(tc):
        if child.tag == qn("w:p"):
            tc.remove(child)


def extract_templates(cell: _Cell) -> dict[str, object]:
    """Extract role-based paragraph prototypes from a template cell.

    Returns a dict with deep-copied <w:p> elements for each role:
      - 'header'        — bold colored header (e.g. "Zakres kontroli (roczna):")
      - 'bullet'        — regular black bullet
      - 'norm_header'   — bold colored "Przepisy / normy / wytyczne:" (only roczny)
      - 'norm_bullet'   — italic gray bullet (only roczny)
      - 'empty'         — empty paragraph (uses bullet style with no text)

    The detection scans cell paragraphs by their text content.
    """
    templates: dict[str, object] = {}
    state = "main"
    for p in cell.paragraphs:
        text = p.text.strip()
        if not text:
            templates.setdefault("empty", deepcopy(p._element))
            continue
        if "header" not in templates and (
            text.startswith("Zakres kontroli")
            or text.startswith("Zakres roczny")
            or text.startswith("Zakres DODATKOWY")
        ):
            templates["header"] = deepcopy(p._element)
            state = "in_zakres"
            continue
        if text.startswith("Przepisy"):
            templates.setdefault("norm_header", deepcopy(p._element))
            state = "in_przepisy"
            continue
        if text.lstrip().startswith("•"):
            if state == "in_przepisy":
                templates.setdefault("norm_bullet", deepcopy(p._element))
            else:
                templates.setdefault("bullet", deepcopy(p._element))
    if "empty" not in templates and "bullet" in templates:
        templates["empty"] = deepcopy(templates["bullet"])
    return templates


def replace_paragraph_text(p_elem, new_text: str) -> None:
    """Replace text inside a <w:p> element while preserving the first run's formatting.

    Removes all but the first run, then replaces its <w:t> with new text.
    """
    runs = p_elem.findall(f"{W_NS}r")
    # Remove all runs beyond the first
    for r in runs[1:]:
        p_elem.remove(r)
    if not runs:
        # Build a fresh run preserving paragraph's existing pPr (if any)
        run = OxmlElement("w:r")
        p_elem.append(run)
    else:
        run = runs[0]
    # Remove all <w:t> and <w:br> from the run, keep <w:rPr>
    for child in list(run):
        if child.tag in (qn("w:t"), qn("w:br")):
            run.remove(child)
    # Add new text element
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = new_text
    run.append(t)


def categorize_line(line: str, state: str) -> tuple[str, str]:
    """Categorize a line and return (role, new_state).

    Roles: 'header', 'bullet', 'norm_header', 'norm_bullet', 'empty'.
    State tracks whether we are inside the 'Przepisy' block (italic norms).
    """
    stripped = line.lstrip()
    if not stripped:
        return "empty", state
    if stripped.startswith("Przepisy"):
        return "norm_header", "in_przepisy"
    is_header_marker = (
        stripped.startswith("Zakres kontroli")
        or stripped.startswith("Zakres roczny")
        or stripped.startswith("Zakres DODATKOWY")
        or stripped.startswith("📋")
        or stripped.startswith("⚙")
        or stripped.startswith("ℹ")
        or stripped.startswith("OBOWIĄZKOWO")
        or stripped.startswith("OPCJONALNIE")
    )
    if is_header_marker:
        return "header", "main"
    # Bullet — within przepisy block uses italic gray (norm_bullet) only for narrowly typical norms
    # We scope norm_bullet only to bullets immediately following 'Przepisy' header
    if stripped.startswith("•"):
        return ("norm_bullet" if state == "in_przepisy" else "bullet"), state
    if stripped.startswith("☐"):
        return "bullet", state
    return "bullet", state


def replace_cell_styled(cell: _Cell, lines: list[str], templates: dict[str, object]) -> None:
    """Replace a cell's content using role-based paragraph templates (deep-copied)."""
    clear_cell(cell)
    state = "main"
    for line in lines:
        role, state = categorize_line(line, state)
        proto = templates.get(role)
        if proto is None:
            proto = templates.get("bullet")
        if proto is None:
            # Fallback: plain paragraph
            p = cell.add_paragraph()
            p.add_run(line)
            continue
        new_p = deepcopy(proto)
        replace_paragraph_text(new_p, line)
        cell._tc.append(new_p)


def insert_paragraph_before(reference_paragraph, text: str, bold: bool = False) -> None:
    """Insert a paragraph immediately before the reference paragraph."""
    new_p = OxmlElement("w:p")
    reference_paragraph._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph

    para = Paragraph(new_p, reference_paragraph._parent)
    run = para.add_run(text)
    run.bold = bold


def find_paragraph_by_text(doc: Document, contains: str):
    """Return first paragraph whose text contains given substring (case-sensitive)."""
    for p in doc.paragraphs:
        if contains in p.text:
            return p
    return None


# ---------- intro note ------------------------------------------------------


INTRO_ROCZNY_R = (
    "UWAGA o zakresie czynności kontrolnych: Niniejszy zakres obejmuje czynności wykonywane w ramach inspekcji "
    "okresowej, tj. kontrolę wizualną, ocenę ekspercką osoby z uprawnieniami budowlanymi oraz weryfikację "
    "dokumentacji. Czynności wymagające specjalistycznego sprzętu lub kompetencji (badania NDT, kontrola momentów "
    "dokręcenia, diagnostyka wibracyjna, dostęp linowy / dronowy do łopat, pomiar luzu łożysk, pomiary grubości "
    "powłok) są realizowane przez certyfikowany serwis techniczny producenta turbiny w ramach umowy serwisowej "
    "(art. 8b ustawy z 20 maja 2016 r. o inwestycjach w zakresie elektrowni wiatrowych). W tabeli ustaleń poniżej "
    "oznaczono je piktogramem ⚙. W przypadku stwierdzenia nieprawidłowości dotyczących takich czynności, w "
    "zaleceniach wskazana jest realizacja przez właściwy serwis."
)

INTRO_ROCZNY_U = (
    "WARIANT UPROSZCZONY (kontrola bez wjazdu na konstrukcję): Inspekcja prowadzona z poziomu terenu i pierwszego "
    "segmentu wieży. Główne narzędzia: oględziny z lornetką, weryfikacja dokumentacji, analiza danych SCADA "
    "udostępnionych przez właściciela / serwis. Wariant odpowiedni dla obiektów objętych pełną umową serwisową "
    "producenta. Czynności wymagające wjazdu (kontrola wyższych segmentów wieży, wnętrza gondoli, wirnika z bliska, "
    "podestów pośrednich, sprzętu BHP/ppoż. w gondoli, kontroli flansz wyższych) realizowane są przez serwis "
    "producenta zgodnie z umową serwisową (art. 8b ustawy z 20 maja 2016 r. o inwestycjach w zakresie elektrowni "
    "wiatrowych)."
)

INTRO_5LETNI = (
    "UWAGA o zakresie czynności kontrolnych: Niniejszy zakres obejmuje czynności wykonywane w ramach inspekcji "
    "5-letniej (z wjazdem na konstrukcję), tj. kontrolę wizualną, ocenę ekspercką, pomiary instalacji elektrycznej "
    "i odgromowej, weryfikację dokumentacji. Czynności wymagające specjalistycznego sprzętu lub kompetencji "
    "(badania NDT, kontrola momentów dokręcenia 100%, diagnostyka wibracyjna, dostęp linowy / dronowy do łopat, "
    "pomiar luzu łożyska wieńcowego, pomiary grubości powłok) są realizowane przez certyfikowany serwis techniczny "
    "producenta turbiny w ramach umowy serwisowej (art. 8b ustawy z 20 maja 2016 r. o inwestycjach w zakresie "
    "elektrowni wiatrowych). W tabeli ustaleń poniżej oznaczono je piktogramem ⚙. Pozycje 📋 są realizowane w "
    "ramach inspekcji 5-letniej."
)

INTRO_5LETNI_O = (
    "WARIANT ODRĘBNY (osobny protokół 5-letni): Niniejszy protokół obejmuje wyłącznie zakres kontroli okresowej "
    "co najmniej raz na 5 lat (art. 62 ust. 1 pkt 2 PB). Kontrola okresowa co najmniej raz w roku (art. 62 ust. 1 "
    "pkt 1 PB) jest sporządzona w odrębnym protokole. Wariant stosowany w przypadkach, gdy właściwy miejscowo "
    "powiatowy inspektorat nadzoru budowlanego wymaga rozdzielenia protokołów. Czynności wymagające specjalistycznego "
    "sprzętu lub kompetencji (badania NDT, kontrola momentów dokręcenia 100%, diagnostyka wibracyjna, dostęp linowy "
    "/ dronowy do łopat, pomiar luzu łożyska wieńcowego, pomiary grubości powłok) są realizowane przez certyfikowany "
    "serwis techniczny producenta turbiny w ramach umowy serwisowej (art. 8b ustawy z 20 maja 2016 r. o inwestycjach "
    "w zakresie elektrowni wiatrowych). W tabeli ustaleń poniżej oznaczono je piktogramem ⚙. Pozycje 📋 są "
    "realizowane w ramach inspekcji 5-letniej."
)

LEGENDA_PIKTOGRAMOW = (
    "Legenda piktogramów: 📋 — w zakresie inspekcji okresowej; "
    "⚙ — wykonuje serwis producenta; "
    "ℹ — realizowane przez podmiot zewnętrzny (np. UDT)."
)


SERWIS_NOTE = (
    "Czynności specjalistyczne oznaczone w tabeli ustaleń piktogramem ⚙ (badania NDT, kontrola momentów dokręcenia, "
    "diagnostyka wibracyjna, inspekcja łopat z dronów / dostępu linowego, pomiar luzu łożysk, próby obciążeniowe "
    "i pomiary grubości powłok) są realizowane przez serwis producenta zgodnie z umową serwisową."
)

# Sekcja VI / IV — Zalecenia (wzór Żeńsko: 6 kolumn + 3 tabele definicji + 3 sekcje narracyjne)
ZALECENIA_HEADERS = [
    "Lp.",
    "Element / lokalizacja",
    "Zakres robót remontowych",
    "Rodzaj",
    "Pilność",
    "Termin wykonania",
]

RODZAJE_ROBOT_HEADERS = ["Symbol", "Rodzaj robót", "Definicja"]
RODZAJE_ROBOT_ROWS = [
    (
        "K",
        "Konserwacja",
        "Roboty utrzymujące sprawność techniczną elementów obiektu (czyszczenie, smarowanie, "
        "dokręcanie, drobne uzupełnienia).",
    ),
    (
        "NB",
        "Naprawa bieżąca",
        "Okresowy remont elementów obiektu zapobiegający skutkom zużycia, utrzymujący właściwy "
        "stan techniczny.",
    ),
    (
        "NG",
        "Naprawa główna",
        "Remont polegający na wymianie co najmniej jednego elementu obiektu.",
    ),
]

STOPNIE_PILNOSCI_HEADERS = ["Stopień", "Zalecany termin", "Opis"]
STOPNIE_PILNOSCI_ROWS = [
    (
        "I",
        "natychmiast",
        "Remont w przypadku uszkodzeń zagrażających bezpieczeństwu użytkowania lub mogących stać "
        "się przyczyną zniszczenia/awarii obiektu. Wymaga natychmiastowego zabezpieczenia, naprawy "
        "głównej, wymiany lub rozbiórki.",
    ),
    (
        "II",
        "do 3 miesięcy",
        "Remont, który może być odłożony na okres do 3 miesięcy lub do okresu zimowego bez szkody "
        "dla użytkowników. Okres przesunięcia winien być wykorzystany na opracowanie dokumentacji "
        "oraz wybór wykonawcy.",
    ),
    (
        "III",
        "do 12 miesięcy",
        "Remont, który może być odłożony na okres do 1 roku bez specjalnej szkody dla użytkowników "
        "obiektu.",
    ),
    (
        "IV",
        "do 5 lat",
        "Remont, który może być odłożony na okres do 5 lat bez specjalnej szkody dla użytkowników "
        "obiektu.",
    ),
]

KRYTERIA_KLASYFIKACJI_HEADERS = ["Klasa", "Zużycie [%]", "Kryterium oceny"]
KRYTERIA_KLASYFIKACJI_ROWS = [
    (
        "dobry",
        "0–15",
        "Element dobrze utrzymany, bez widocznego zużycia. Cechy materiałów odpowiadają wymogom "
        "norm. Ewentualne drobne naprawy konserwacyjne.",
    ),
    (
        "zadowalający",
        "16–30",
        "Element utrzymywany należycie. Celowe wykonanie konserwacji lub napraw bieżących w "
        "niewielkim zakresie.",
    ),
    (
        "średni",
        "31–50",
        "Niewielkie uszkodzenia i ubytki, niezagrażające bezpieczeństwu. Wymagana naprawa bieżąca "
        "w większym zakresie lub naprawa główna.",
    ),
    (
        "zły",
        "51–70",
        "Znaczne ubytki mogące zagrażać bezpieczeństwu. Materiały utraciły pierwotne właściwości. "
        "Wymagany remont kapitalny — wymiana wielu elementów.",
    ),
    (
        "awaryjny",
        "> 71",
        "Tak duże zniszczenia/ubytki, że nie pozwalają na dalsze bezpieczne użytkowanie. Wymagany "
        "remont kapitalny w dużym rozmiarze lub rozbiórka.",
    ),
]

DEFINICJE_ROBOT_HEADER = "Definicje rodzajów robót remontowych"
STOPNIE_PILNOSCI_HEADER = "Zalecany czas wykonania robót remontowych — stopień pilności"
KRYTERIA_KLASYFIKACJI_HEADER = "Kryteria oceny i klasyfikacji stanu technicznego"

# Sekcje narracyjne (po Żeńsko, dla 5-letni i rocznych)
NARRATYWNE_SECTIONS = [
    (
        "Stan techniczny instalacji ochrony środowiska",
        "(opisz wynik przeglądu instalacji i urządzeń służących ochronie środowiska — instalacja "
        "odgromowa, oświetlenie nawigacyjne, oddziaływania środowiskowe)",
    ),
    (
        "Weryfikacja kompletności i aktualności dokumentów",
        "(opisz wynik weryfikacji KOB, protokołów serwisowych, protokołów pomiarowych, certyfikatów "
        "UDT urządzeń podlegających kontroli)",
    ),
    (
        "Metody i środki użytkowania elementów narażonych na szkodliwe wpływy atmosferyczne i "
        "niszczące działanie innych czynników",
        "(opisz zastosowane metody / środki ochrony lub wpisz „Nie dotyczy”)",
    ),
]

# Sekcja VII — siatka fotografii (wzór Żeńsko)
SEKCJA_VII_SUBTITLE = (
    "Dokumentacja fotograficzna wykonana podczas kontroli (elementy obiektu posiadające usterki "
    "lub wady, przewidziane do remontu)"
)
PHOTO_GRID_ROWS = 6
PHOTO_GRID_COLS = 2  # 6 × 2 = 12 placeholderów

# Sekcja IV — wzór Żeńsko (uproszczona forma: podsumowanie + załącznik PDF z pomiarami)
POMIAR_PODSUMOWANIE_HEADER = "Podsumowanie pomiarów"
POMIAR_PODSUMOWANIE_ROWS = [
    ("Nr protokołu z pomiaru", ""),
    ("Data pomiaru", ""),
    ("Data kolejnego pomiaru", ""),
    ("Orzeczenie", "Dopuszcza / Warunkowo dopuszcza / Nie dopuszcza"),
    ("Ocena końcowa", "☐ Pozytywna   ☐ Negatywna"),
    ("Oględziny instalacji elektrycznej", "☐ Pozytywna   ☐ Negatywna"),
    ("Oględziny instalacji odgromowej i uziomów", "☐ Pozytywna   ☐ Negatywna"),
    ("Uwagi do oględzin i oceny", ""),
]
POMIAR_ZALACZNIK_NOTE = "Pełny protokół pomiarów stanowi załącznik do niniejszej kontroli (PDF)."
POMIAR_PRZYRZADY_HEADER = "Identyfikacja użytych przyrządów"
POMIAR_PRZYRZADY_COLS = ["Model", "Numer seryjny", "Producent"]
POMIAR_OSOBY_HEADER = "Osoby wykonujące pomiary"
POMIAR_OSOBY_COLS = ["Imię i nazwisko", "Numer uprawnień", "Izba"]


# ---------- ROCZNY R (rozszerzona) ------------------------------------------
# Tylko 3 elementy mają zmiany w C1 vs oryginał: E3, E6, E7

ROCZNY_R_E3_C1 = [
    "Zakres kontroli (roczna):",
    "• Wizualna kontrola zewnętrzna (z dołu, z użyciem lornetki) i wewnętrzna wieży po wjeździe",
    "• Sprawdzenie szczelności wpustów kablowych",
    "• Ocena stanu drabin wewnętrznych i mocowań",
    "Przepisy / normy / wytyczne:",
    "• art. 62 ust. 1 pkt 1 lit. a i pkt 2 PB",
    "• PN-EN 1993 (Eurokod 3)",
    "• PN-EN ISO 12944 – ochrona powłokami malarskimi",
    "• PN-EN 50308 – eksploatacja i utrzymanie ruchu",
]

ROCZNY_R_E6_C1 = [
    "Zakres kontroli (roczna):",
    "• Ocena wizualna wirnika i łopat z poziomu terenu (lornetka) oraz z gondoli",
    "• Oględziny receptorów odgromowych w dostępnym zakresie (z gondoli)",
    "• Analiza alarmów i ostrzeżeń ze SCADA z ostatnich 12 miesięcy (jeżeli udostępniono dostęp do danych)",
    "⚙ Inspekcja łopat z bliska (drony, dostęp linowy) — serwis producenta",
    "Przepisy / normy / wytyczne:",
    "• art. 8b ustawy z 20.05.2016 r. o inwestycjach w zakresie elektrowni wiatrowych",
    "• PN-EN IEC 61400-1 – wymagania projektowe",
    "• PN-EN IEC 61400-23 – pełne badanie strukturalne łopat",
    "• PN-EN 62305-3 – ochrona odgromowa łopat",
]

ROCZNY_R_E7_C1 = [
    "Zakres kontroli (roczna):",
    "• Wizualna kontrola połączeń śrubowych piasta–wał (oznaczenia momentu, korozja)",
    "• Oględziny łożyska głównego — wycieki smaru, nietypowe dźwięki podczas pracy",
    "• Analiza wskazań temperatury i drgań ze SCADA (jeżeli udostępniono dostęp)",
    "⚙ Diagnostyka wibracyjna łożyska, kontrola momentów piasta–wał — serwis producenta",
    "Przepisy / normy / wytyczne:",
    "• PN-EN 1090-2 – połączenia sprężane",
    "• PN-EN 14399 – zestawy śrubowe sprężane",
    "• PN-EN IEC 61400-4 – łożyska i przekładnie",
    "• Instrukcja serwisowa producenta turbiny",
]


# 5-LETNI — kolumna C1 (Zakres roczny poszerzony) — analogiczne zmiany jak w R E3/E6/E7
# (5-letni C1 nie zawiera „Przepisy / normy" — to jest w C0)
P5_E3_C1 = [
    "Zakres roczny (oględziny):",
    "• Wizualna kontrola zewnętrzna (z dołu, z użyciem lornetki) i wewnętrzna wieży po wjeździe",
    "• Sprawdzenie szczelności wpustów kablowych",
    "• Ocena stanu drabin wewnętrznych i mocowań",
]

P5_E6_C1 = [
    "Zakres roczny (oględziny):",
    "• Ocena wizualna wirnika i łopat z poziomu terenu (lornetka) oraz z gondoli",
    "• Oględziny receptorów odgromowych w dostępnym zakresie (z gondoli)",
    "• Analiza alarmów i ostrzeżeń ze SCADA z ostatnich 12 miesięcy (jeżeli udostępniono dostęp do danych)",
    "⚙ Inspekcja łopat z bliska (drony, dostęp linowy) — serwis producenta",
]

P5_E7_C1 = [
    "Zakres roczny (oględziny):",
    "• Wizualna kontrola połączeń śrubowych piasta–wał (oznaczenia momentu, korozja)",
    "• Oględziny łożyska głównego — wycieki smaru, nietypowe dźwięki podczas pracy",
    "• Analiza wskazań temperatury i drgań ze SCADA (jeżeli udostępniono dostęp)",
    "⚙ Diagnostyka wibracyjna łożyska, kontrola momentów piasta–wał — serwis producenta",
]


# ---------- ROCZNY U (uproszczona) ------------------------------------------
# Tabela #6 ma 31 wierszy: row 0 = nagłówek, row 1 = E1 nazwa,
# row 2 = E1 dane, row 3 = E2 nazwa, row 4 = E2 dane, ... (parzyste = nazwa, nieparzyste = dane)
# Zmieniamy C1 (OPIS STANU TECHNICZNEGO) dla każdego z 15 elementów.
# C0 (Pozycje do oceny + Przepisy) zostawiamy bez zmian — to nazwy do oceny, nie sam zakres czynności.

ROCZNY_U = {
    1: [  # FUNDAMENT — bez zmian
        "Zakres kontroli (roczna):",
        "• Ocena wizualna stanu fundamentu (bez odkrywek)",
        "• Pomiar/oględziny szerokości rys",
        "• Sprawdzenie zabezpieczenia antykorozyjnego kotew widocznych",
        "• Sprawdzenie szczelności styku fundament–segment 0 wieży",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 1 lit. a i pkt 2 PB",
        "• PN-EN 1992 (Eurokod 2) – konstrukcje betonowe",
        "• PN-EN 1997 (Eurokod 7) – geotechnika",
        "• Wytyczne producenta turbiny",
    ],
    2: [  # FLANSZE
        "Zakres kontroli (roczna):",
        "• Wizualna ocena flanszy dolnej fundament–segment 1 (z poziomu terenu)",
        "• Wizualna kontrola flanszy w obrębie pierwszego segmentu wieży (od wewnątrz, z parteru)",
        "• Ocena stanu powłok antykorozyjnych w dostępnym zakresie",
        "Przepisy / normy / wytyczne:",
        "• PN-EN 1090-2 – konstrukcje stalowe (klasa EXC3/EXC4)",
        "• PN-EN 14399 – zestawy śrubowe sprężane",
        "• PN-EN 1993 (Eurokod 3) – konstrukcje stalowe",
        "• Instrukcja serwisowa producenta turbiny",
    ],
    3: [  # WIEŻA
        "Zakres kontroli (roczna):",
        "• Wizualna kontrola zewnętrzna z poziomu terenu (lornetka, kilka punktów obserwacyjnych dookoła turbiny)",
        "• Ocena wewnętrzna pierwszego segmentu wieży (parter — drabina dolna, oświetlenie, wpust kablowy, uszczelnienia, wnęki)",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 1 lit. a i pkt 2 PB",
        "• PN-EN 1993 (Eurokod 3)",
        "• PN-EN ISO 12944 – ochrona powłokami malarskimi",
        "• PN-EN 50308 – eksploatacja i utrzymanie ruchu",
    ],
    4: [  # GONDOLA
        "Zakres kontroli (roczna):",
        "• Ocena wizualna gondoli z poziomu terenu (lornetka; w miarę możliwości obrót gondoli o 360° dla pełnego oglądu)",
        "• Brak widocznych wycieków, zacieków na zewnętrznej obudowie",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 1 lit. a i pkt 2 PB",
        "• PN-EN 50308 – eksploatacja i utrzymanie ruchu",
        "• PN-EN IEC 61400-1 – wymagania projektowe",
        "• Wytyczne producenta turbiny",
    ],
    5: [  # YAW
        "Zakres kontroli (roczna):",
        "• Nasłuch podczas obrotu gondoli (yaw test) — brak nietypowych dźwięków, drgań przekazywanych na konstrukcję",
        "• Brak widocznych wycieków smaru / oleju na zewnętrznej powierzchni gondoli i na ramie",
        "• Analiza alarmów ze SCADA dot. systemu yaw (jeżeli udostępniono)",
        "Przepisy / normy / wytyczne:",
        "• PN-EN 1090-2 – połączenia sprężane",
        "• PN-EN 14399 – zestawy śrubowe sprężane",
        "• PN-EN IEC 61400-4 – łożyska i przekładnie turbiny",
        "• Instrukcja serwisowa producenta turbiny",
    ],
    6: [  # WIRNIK
        "Zakres kontroli (roczna):",
        "• Ocena wizualna wirnika i łopat z poziomu terenu (lornetka)",
        "• Wizualna ocena receptorów odgromowych z dołu w dostępnym zakresie",
        "• Analiza alarmów i ostrzeżeń ze SCADA dot. wirnika i pitch (jeżeli udostępniono)",
        "Przepisy / normy / wytyczne:",
        "• art. 8b ustawy z 20.05.2016 r. o inwestycjach w zakresie elektrowni wiatrowych",
        "• PN-EN IEC 61400-1 – wymagania projektowe",
        "• PN-EN IEC 61400-23 – pełne badanie strukturalne łopat",
        "• PN-EN 62305-3 – ochrona odgromowa łopat",
    ],
    7: [  # PIASTA / ŁOŻYSKO GŁÓWNE / PITCH
        "Zakres kontroli (roczna):",
        "• Analiza wskazań temperatury i drgań ze SCADA (jeżeli udostępniono dostęp do danych)",
        "• Brak widocznych wycieków smaru / oleju na zewnętrznej obudowie gondoli / pod gondolą",
        "Przepisy / normy / wytyczne:",
        "• PN-EN 1090-2 – połączenia sprężane",
        "• PN-EN 14399 – zestawy śrubowe sprężane",
        "• PN-EN IEC 61400-4 – łożyska i przekładnie",
        "• Instrukcja serwisowa producenta turbiny",
    ],
    8: [  # PODESTY
        "Zakres kontroli (roczna):",
        "• Stan techniczny podestu wejściowego (parter wieży) — oględziny, mocowania, korozja",
        "• Stan dolnej części drabiny i jej mocowań do segmentu 1",
        "• Sprawdzenie sprawności oświetlenia awaryjnego w pierwszym segmencie",
        "Przepisy / normy / wytyczne:",
        "• PN-EN 50308 – zabezpieczenia, eksploatacja",
        "• PN-EN ISO 14122 – stałe środki dostępu do maszyn",
        "• PN-EN 353-1 – urządzenia samohamowne (rail)",
        "• Wytyczne producenta",
    ],
    9: [  # EWAKUACJA
        "Zakres kontroli (roczna):",
        "• Kontrola sprzętu ewakuacyjno-ratunkowego znajdującego się na poziomie terenu / pierwszego segmentu (kompletność, daty ważności)",
        "• Apteczka pierwszej pomocy — stan, kompletność",
        "• Sprawdzenie czytelności i aktualności planu ewakuacji",
        "Przepisy / normy / wytyczne:",
        "• PN-EN 50308 – wymagania dot. ewakuacji",
        "• PN-EN 341 – urządzenia do opuszczania",
        "• Rozp. MGiP w sprawie BHP w energetyce",
        "• Wytyczne producenta",
    ],
    10: [  # SCHODY ZEWNĘTRZNE — bez zmian
        "Zakres kontroli (roczna):",
        "• Ocena schodów zewnętrznych (oględziny, korozja, ugięcia)",
        "• Sprawdzenie szczelności i mocowań drzwi zewnętrznych",
        "• Ocena stanu balustrad / pochwytów",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 1 lit. a i pkt 2 PB",
        "• PN-EN 50308",
        "• PN-EN ISO 14122 – stałe środki dostępu",
        "• Wytyczne producenta",
    ],
    11: [  # LPS
        "Zakres kontroli (roczna):",
        "• Wizualna kontrola dolnych elementów instalacji odgromowej (uziomy widoczne, złącza kontrolne na fundamencie / segmencie 1)",
        "• Sprawdzenie integralności złącz kontrolnych w dostępnym zakresie",
        "• Oględziny wskaźników stanu SPD na rozdzielnicy parteru",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 2 PB (obowiązkowo co 5 lat)",
        "• PN-EN 62305 (seria) – ochrona odgromowa",
        "• PN-EN 62305-3 – ochrona obiektów (LPS)",
        "• PN-HD 60364-5-54 – uziemienie",
    ],
    12: [  # INST. ELEKTRYCZNA
        "Zakres kontroli (roczna):",
        "• Oględziny rozdzielnicy parteru / stacji rozdzielczej (oznaczenia, dostępność, kompletność osłon)",
        "• Oględziny transformatora w stacji rozdzielczej (poziom oleju, brak wycieków, brak nietypowych dźwięków)",
        "• Stan wpustu kablowego SN / nN i kabli na parterze wieży",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 2 PB (obowiązkowo co 5 lat)",
        "• Ustawa Prawo energetyczne",
        "• Rozp. ME – BHP przy urządzeniach energetycznych",
        "• PN-HD 60364 – instalacje elektryczne niskiego napięcia",
    ],
    13: [  # BHP / PPOŻ
        "Zakres kontroli (roczna):",
        "• Termin ważności przeglądów gaśnic na parterze i w stacji rozdzielczej",
        "• Sprawdzenie sprawności czujek dymowych (jeżeli zainstalowane na parterze)",
        "• Oględziny oznakowania ewakuacyjnego dolnej części wieży",
        "Przepisy / normy / wytyczne:",
        "• Rozp. MSWiA – ochrona ppoż. budynków",
        "• PN-EN 3 (seria) – gaśnice przenośne",
        "• PN-EN 54 – systemy sygnalizacji pożarowej",
        "• Ustawa o ochronie ppoż.",
    ],
    14: [  # UDT — bez zmian (dokumentacyjnie)
        "Zakres kontroli (roczna):",
        "• Aktualność badań UDT i protokołów odbiorów",
        "• Stan techniczny tabliczek znamionowych i znaków UDT (urządzenia dostępne)",
        "• Aktualność szkoleń operatorów",
        "Przepisy / normy / wytyczne:",
        "• Ustawa o dozorze technicznym",
        "• Rozp. RM w sprawie rodzajów urządzeń technicznych podlegających dozorowi",
        "• Rozp. MR w sprawie warunków technicznych UDT",
        "• DTR producenta",
    ],
    15: [  # DOJAZDY — bez zmian
        "Zakres kontroli (roczna):",
        "• Ocena nawierzchni drogi i placu (ubytki, zagłębienia, koleiny)",
        "• Sprawdzenie drożności odwodnień",
        "• Oględziny ogrodzenia / oznakowania",
        "Przepisy / normy / wytyczne:",
        "• art. 62 ust. 1 pkt 1 lit. a i b oraz pkt 2 PB",
        "• Rozp. MTBiGM – warunki techniczne dróg",
        "• PN-EN 13242 – kruszywa drogowe",
        "• Wytyczne producenta",
    ],
}


# ---------- 5-LETNI ---------------------------------------------------------
# Tabela #8 ma 33 wiersze: row 0 = nagłówek (z 7 kolumnami),
# row 1 = E1 nazwa, row 2 = E1 dane, ...
# Zmieniamy kolumnę C2 (ZAKRES DODATKOWY 5-LETNI) dla każdego z 16 elementów.
# Niektóre elementy (E11 LPS, E12 elektryka) zostają z dodanym checkboxem termowizji w E12.

P5_C2 = {
    1: [  # FUNDAMENT
        "📋 W zakresie inspekcji 5-letniej:",
        "• Pełna ocena stanu technicznego i przydatności do użytkowania",
        "• Wizualna ocena izolacji przeciwwilgociowej w dostępnych miejscach",
    ],
    2: [  # FLANSZE
        "📋 W zakresie inspekcji 5-letniej:",
        "• Szczegółowa wizualna inspekcja flansz i powłok ochronnych po wjeździe",
        "• Weryfikacja aktualności oznaczeń kontrolnych momentu",
        "",
        "⚙ Wykonuje serwis producenta (art. 8b u.w.):",
        "• Pełna kontrola momentów dokręcenia (100% lub statystycznie wg producenta)",
        "• Wymiana / kontrola śrub o znamionach pęknięć / korozji wżerowej",
        "• Badania NDT spawów flansz (UT/MT) wg wytycznych producenta",
    ],
    3: [  # WIEŻA
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowa wizualna ocena stanu i przydatności do użytkowania",
        "• Ocena powłok antykorozyjnych na całej dostępnej powierzchni",
    ],
    4: [  # GONDOLA
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowa wizualna ocena stanu konstrukcji gondoli",
        "• Ocena instalacji wewnątrz gondoli (oświetlenie, gniazda, wentylacja)",
    ],
    5: [  # YAW
        "📋 W zakresie inspekcji 5-letniej:",
        "• Wizualna ocena uzębienia wieńca i napędów azymutu (zakres dostępny z gondoli)",
        "",
        "⚙ Wykonuje serwis producenta (art. 8b u.w.):",
        "• Pełna kontrola momentów dokręcenia śrub w połączeniu wieża–łożysko",
        "• Pomiar luzu łożyska wieńcowego wg instrukcji producenta",
        "• Pełna ocena stanu uzębienia wieńca i hamulców azymutu",
    ],
    6: [  # WIRNIK
        "📋 W zakresie inspekcji 5-letniej:",
        "• Wizualna ocena łopat z gondoli (z bliska — strona nawietrzna i zawietrzna) oraz z poziomu terenu (lornetka)",
        "• Wizualna ocena receptorów odgromowych w dostępnym zakresie",
        "",
        "⚙ Wykonuje serwis producenta (art. 8b u.w.):",
        "• Szczegółowa inspekcja łopat z bliska (drony, kamery termowizyjne, dostęp linowy)",
        "• Pomiar ciągłości obwodu odgromowego łopat",
        "• Pełna kontrola momentów dokręcenia połączeń łopata–piasta",
        "• Diagnostyka systemu pitch (łożyska, napędy, akumulatory awaryjne)",
    ],
    7: [  # PIASTA
        "📋 W zakresie inspekcji 5-letniej:",
        "• Wizualna ocena uszczelnień, układu centralnego smarowania (z gondoli)",
        "",
        "⚙ Wykonuje serwis producenta (art. 8b u.w.):",
        "• Pełna kontrola momentów dokręcenia połączeń śrubowych",
        "• Diagnostyka wibracyjna łożyska głównego (analizator drgań)",
        "• Pełna kontrola układu pitch",
    ],
    8: [  # PODESTY
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowa wizualna ocena nośności i przydatności podestów",
        "• Pełen wizualny przegląd systemu asekuracji wzdłuż drabiny (rail/wire — korozja, mocowania, kompletność)",
    ],
    9: [  # EWAKUACJA
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowy wizualny przegląd zestawu ewakuacyjnego — kompletność, daty ważności, stan techniczny",
        "• Weryfikacja aktualności planu ewakuacji",
    ],
    10: [  # SCHODY ZEW.
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowa wizualna ocena stanu technicznego i przydatności",
        "• Pełen przegląd systemów zamknięć drzwi (zamki, uszczelki, sprężyny)",
    ],
    11: [  # LPS
        "📋 W zakresie inspekcji 5-letniej (OBOWIĄZKOWO — art. 62 ust. 1 pkt 2 PB):",
        "• Pomiar rezystancji uziemienia wszystkich uziomów",
        "• Pomiar ciągłości rezystancji LPS (zwody, przewody odprowadzające, złącza kontrolne)",
        "• Sprawdzenie sprawności ograniczników przepięć (SPD)",
        "• Sporządzenie protokołu pomiarów (osoba z uprawnieniami pomiarowymi E/D + SEP do 1 kV)",
    ],
    12: [  # INST. ELEKTRYCZNA — z checkboxem termowizji
        "📋 W zakresie inspekcji 5-letniej:",
        "",
        "OBOWIĄZKOWO (art. 62 ust. 1 pkt 2 PB):",
        "• Pomiar rezystancji izolacji obwodów",
        "• Pomiar pętli zwarcia obwodów gniazd 230 V",
        "• Sprawdzenie ciągłości przewodów ochronnych",
        "• Sprawdzenie sprawności wyłączników różnicowo-prądowych (RCD)",
        "",
        "OPCJONALNIE (zaznaczyć po wykonaniu):",
        "☐ Termowizja rozdzielnic i połączeń kablowych (przy nominalnym obciążeniu turbiny)",
    ],
    13: [  # BHP / PPOŻ
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowy wizualny przegląd wszystkich elementów ppoż. (gaśnice, czujki, oznakowanie)",
        "• Weryfikacja aktualności instrukcji bezpieczeństwa pożarowego (IBP)",
    ],
    14: [  # UDT
        "📋 W zakresie inspekcji 5-letniej:",
        "• Przegląd resursu urządzeń UDT na podstawie dokumentacji",
        "• Ocena przydatności do dalszej eksploatacji",
        "• Weryfikacja kompletności dokumentacji DTR / instrukcji",
        "",
        "ℹ Badania techniczne UDT (okresowe / doraźne) realizuje Urząd Dozoru Technicznego w terminach z decyzji UDT.",
    ],
    15: [  # DOJAZDY
        "📋 W zakresie inspekcji 5-letniej:",
        "• Kompleksowa wizualna ocena stanu technicznego nawierzchni i odwodnień",
        "• Weryfikacja aktualności oznakowania ostrzegawczego",
    ],
    16: [  # ESTETYKA — bez zmian (zostaje pełen zakres)
        "Zakres DODATKOWY 5-letni:",
        "• Ocena estetyki obiektu na podstawie wizji lokalnej (OBOWIĄZKOWO)",
        "• Ocena estetyki otoczenia obiektu (drogi, plac, ogrodzenie, zieleń)",
        "• Ocena oznakowania ostrzegawczego (lampy nocne, pasy malowania)",
        "• Ujęcie potrzebnych prac estetycznych w planie remontów",
    ],
}


# ---------- generation funcs ------------------------------------------------


INTRO_MARKERS = (
    "UWAGA o zakresie czynności kontrolnych",
    "WARIANT UPROSZCZONY",
    "Legenda piktogramów:",
)
SERWIS_MARKER = "Czynności specjalistyczne oznaczone w tabeli ustaleń piktogramem"


def remove_paragraphs_with_markers(doc: Document, markers: tuple[str, ...]) -> None:
    """Idempotency helper — remove any existing paragraph that matches a marker substring."""
    for p in list(doc.paragraphs):
        if any(m in p.text for m in markers):
            p._element.getparent().remove(p._element)


def add_intro_above_table_iii(
    doc: Document,
    intro_text: str,
    *,
    include_legenda: bool = True,
    extra_text: str | None = None,
) -> None:
    """Insert intro paragraphs immediately before 'W trakcie kontroli ustalono' or table III header.

    Idempotent — clears any existing intro paragraphs before adding.
    """
    remove_paragraphs_with_markers(doc, INTRO_MARKERS)

    target = find_paragraph_by_text(doc, "W trakcie kontroli ustalono")
    if target is None:
        target = find_paragraph_by_text(doc, "III. USTALENIA")
    if target is None:
        raise RuntimeError("Could not find anchor paragraph for intro note")

    insert_paragraph_before(target, intro_text, bold=False)
    if include_legenda:
        insert_paragraph_before(target, LEGENDA_PIKTOGRAMOW, bold=False)
    if extra_text:
        insert_paragraph_before(target, extra_text, bold=False)


def remove_serwis_checklist(doc: Document) -> None:
    """Remove 'Zakres czynności serwisowych (zaznaczyć wykonane)' header and the
    following checkbox table (per user request 2026-05-11)."""
    target = find_paragraph_by_text(doc, "Zakres czynności serwisowych")
    if target is None:
        return  # Already removed
    next_elem = target._element.getnext()
    target._element.getparent().remove(target._element)
    if next_elem is not None and next_elem.tag == qn("w:tbl"):
        next_elem.getparent().remove(next_elem)


def add_serwis_note(doc: Document) -> None:
    """Insert serwis note paragraph at the end of section V (before next 'VI. ...' header).

    Idempotent — clears any existing serwis note before adding.
    """
    remove_paragraphs_with_markers(doc, (SERWIS_MARKER,))

    v_para = find_paragraph_by_text(doc, "V. INFORMACJE O SERWISIE")
    if v_para is None:
        return

    # Walk forward from V. INFORMACJE to find next paragraph starting with "VI."
    elem = v_para._element.getnext()
    while elem is not None:
        if elem.tag == qn("w:p"):
            text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
            if text.strip().startswith("VI."):
                new_p = _make_paragraph_element(SERWIS_NOTE, bold=False)
                elem.addprevious(new_p)
                return
        elem = elem.getnext()


def update_table_iii_description_for_O(doc: Document) -> None:
    """Update the italic paragraph above tabela III that describes its columns.

    In combined 5-letni it references both 'Zakres roczny' and 'Zakres dodatkowy 5-letni'.
    In _O variant only 'Zakres 5-letni' column exists — so we shorten the description.
    """
    target = find_paragraph_by_text(doc, "W trakcie kontroli ustalono. Kolumna")
    if target is None:
        return
    new_text = (
        "W trakcie kontroli ustalono. Kolumna „Zakres 5-letni” wskazuje "
        "czynności obowiązkowe co 5 lat (art. 62 ust. 1 pkt 2 PB)."
    )
    runs = target.runs
    if runs:
        runs[0].text = new_text
        for r in runs[1:]:
            r._element.getparent().remove(r._element)
    else:
        target.add_run(new_text)


def remove_uwaga_combined_scope(doc: Document) -> None:
    """Remove the standalone 1x1 UWAGA table about full annual scope inclusion (5-letni only)."""
    for tbl in list(doc.tables):
        if len(tbl.rows) == 1 and len(tbl.columns) == 1:
            text = tbl.cell(0, 0).text
            if "Kontrola pięcioletnia obejmuje pełny zakres" in text:
                tbl._element.getparent().remove(tbl._element)
                return


def remove_column_from_table(table, col_index: int) -> None:
    """Remove a column from a table, handling gridSpan-merged cells.

    Updates <w:tblGrid>, removes <w:tc> elements at the column position, or
    decrements gridSpan when a cell spans across the removed column.
    """
    tbl = table._tbl

    # Update <w:tblGrid>
    grid = tbl.find(f"{W_NS}tblGrid")
    if grid is not None:
        cols = grid.findall(f"{W_NS}gridCol")
        if col_index < len(cols):
            grid.remove(cols[col_index])

    # Update each row
    for row in tbl.findall(f"{W_NS}tr"):
        tcs = row.findall(f"{W_NS}tc")
        current_col = 0
        for tc in tcs:
            tcPr = tc.find(f"{W_NS}tcPr")
            gridSpan_elem = tcPr.find(f"{W_NS}gridSpan") if tcPr is not None else None
            span = int(gridSpan_elem.get(qn("w:val"))) if gridSpan_elem is not None else 1
            if current_col + span > col_index:
                # This cell covers col_index
                if span == 1:
                    row.remove(tc)
                else:
                    new_span = span - 1
                    if new_span == 1:
                        tcPr.remove(gridSpan_elem)
                    else:
                        gridSpan_elem.set(qn("w:val"), str(new_span))
                break
            current_col += span


def apply_pomiar_template(doc: Document) -> None:
    """Replace section IV (Wyniki pomiarów elektrycznych) with the simplified Żeńsko-style template.

    Removes:
      - Existing A/B/C subsection labels and tables (instalacja elektryczna, odgromowa, wykaz protokołów)
      - 'WAŻNE: Instalacja LPS...' paragraph

    Inserts (after 'Pomiary wykonano zgodnie z...' paragraph):
      - 'Podsumowanie pomiarów' (header + 2-col table, 8 rows)
      - Note 'Pełny protokół pomiarów stanowi załącznik...'
      - 'Identyfikacja użytych przyrządów' (header + 3-col table, 4 rows incl. header)
      - 'Osoby wykonujące pomiary' (header + 3-col table, 4 rows incl. header)

    Idempotent — detects if Żeńsko template is already applied (via Podsumowanie header)
    and skips re-applying.
    """
    # Idempotency check
    if find_paragraph_by_text(doc, POMIAR_PODSUMOWANIE_HEADER) is not None:
        return  # Already applied

    # Find boundaries: insert after intro, remove until V. INFORMACJE
    intro = find_paragraph_by_text(doc, "Pomiary wykonano zgodnie z art. 62")
    section_v = find_paragraph_by_text(doc, "V. INFORMACJE O SERWISIE")
    if intro is None or section_v is None:
        raise RuntimeError("Could not find section IV intro or section V anchor")

    # Collect all elements between intro (exclusive) and section_v (exclusive)
    to_remove = []
    elem = intro._element.getnext()
    while elem is not None and elem != section_v._element:
        to_remove.append(elem)
        elem = elem.getnext()
    for elem in to_remove:
        elem.getparent().remove(elem)

    # Now insert new structure BEFORE section_v
    # 1. Podsumowanie pomiarów header
    insert_paragraph_before(section_v, POMIAR_PODSUMOWANIE_HEADER, bold=True)
    # 2. Podsumowanie pomiarów table (2 cols × 8 rows) — label/value, no header row,
    # so use borders-only styling + bold first column for labels.
    table_pods = _insert_table_before(doc, section_v, rows=len(POMIAR_PODSUMOWANIE_ROWS), cols=2)
    for ri, (label, value) in enumerate(POMIAR_PODSUMOWANIE_ROWS):
        _set_cell_text_bold(table_pods.cell(ri, 0), label, bold=True)
        table_pods.cell(ri, 1).text = value
    _style_borders_only(table_pods)
    # 3. Załącznik note
    insert_paragraph_before(section_v, POMIAR_ZALACZNIK_NOTE, bold=False)
    # 4. Identyfikacja przyrządów header
    insert_paragraph_before(section_v, POMIAR_PRZYRZADY_HEADER, bold=True)
    # 5. Identyfikacja przyrządów table (3 cols × 4 rows: 1 header + 3 data)
    table_przyrz = _insert_table_before(doc, section_v, rows=4, cols=3)
    for ci, col_name in enumerate(POMIAR_PRZYRZADY_COLS):
        _set_cell_text_bold(table_przyrz.cell(0, ci), col_name, bold=True)
    style_table_zensko(table_przyrz)
    # 6. Osoby wykonujące header
    insert_paragraph_before(section_v, POMIAR_OSOBY_HEADER, bold=True)
    # 7. Osoby wykonujące table (3 cols × 4 rows: 1 header + 3 data)
    table_osoby = _insert_table_before(doc, section_v, rows=4, cols=3)
    for ci, col_name in enumerate(POMIAR_OSOBY_COLS):
        _set_cell_text_bold(table_osoby.cell(0, ci), col_name, bold=True)
    style_table_zensko(table_osoby)


def _insert_table_before(doc: Document, reference_para, rows: int, cols: int):
    """Create a new table and move its XML element before the reference paragraph.

    Uses 'Table Grid' style for borders.
    """
    table = doc.add_table(rows=rows, cols=cols)
    try:
        table.style = doc.styles["Table Grid"]
    except KeyError:
        pass
    # Move table XML element from end-of-doc to before reference
    reference_para._element.addprevious(table._element)
    return table


def style_table_zensko(table, zebra: bool = True) -> None:
    """Apply Żeńsko-style styling: thin grey borders, dark navy header with white bold text,
    optional zebra striping on body rows (light grey on every second row).
    """
    tbl = table._tbl

    # Find or create tblPr
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    # Replace tblBorders with thin grey
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), TABLE_BORDER_COLOR)
        tblBorders.append(b)
    tblPr.append(tblBorders)

    # Style header row (row 0): dark navy fill + white bold text
    for cell in table.rows[0].cells:
        tcPr = cell._tc.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = OxmlElement("w:tcPr")
            cell._tc.insert(0, tcPr)
        existing_shd = tcPr.find(qn("w:shd"))
        if existing_shd is not None:
            tcPr.remove(existing_shd)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), TABLE_HEADER_FILL)
        tcPr.append(shd)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = TABLE_HEADER_TEXT

    # Optional zebra striping on body rows
    if zebra:
        for ri, row in enumerate(table.rows[1:], start=1):
            if ri % 2 == 0:
                continue  # leave white
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcPr)
                existing_shd = tcPr.find(qn("w:shd"))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), TABLE_ZEBRA_FILL)
                tcPr.append(shd)


def fix_section_header_rows(doc: Document) -> None:
    """Nadaje granatowy fill TABLE_HEADER_FILL wierszom „section header" w tabeli III.

    Wiersze typu „1. FUNDAMENT", „6. WIRNIK / ROTOR Z ŁOPATAMI" itp. są scalone
    przez całą szerokość tabeli (gridSpan == num_cols) i mają biały bold tekst,
    ale źródłowy docx ma na nich jasnoszary fill `F5F7F9` (zebra) — co czyni
    tekst nieczytelnym. Ta funkcja zamienia fill na granatowy `1B2230` i utrwala
    bold + biały kolor tekstu.
    """
    for table in doc.tables:
        if not table.rows:
            continue
        num_cols = len(table.columns)
        if num_cols < 2:
            continue
        for row in table.rows[1:]:  # row 0 to nagłówek kolumn — pomijamy
            if not row.cells:
                continue
            first_tcPr = row.cells[0]._tc.find(qn("w:tcPr"))
            gs_el = first_tcPr.find(qn("w:gridSpan")) if first_tcPr is not None else None
            gs_val = int(gs_el.get(qn("w:val"))) if gs_el is not None else 1
            if gs_val < num_cols:
                continue
            for cell in row.cells:
                tcPr = cell._tc.find(qn("w:tcPr"))
                if tcPr is None:
                    tcPr = OxmlElement("w:tcPr")
                    cell._tc.insert(0, tcPr)
                existing_shd = tcPr.find(qn("w:shd"))
                if existing_shd is not None:
                    tcPr.remove(existing_shd)
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), TABLE_HEADER_FILL)
                tcPr.append(shd)
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = TABLE_HEADER_TEXT


def apply_arial_font_globally(doc: Document) -> None:
    """Change font of all runs (paragraphs + table cells) to Arial.

    Preserves bold/italic/color/size from original — only replaces font name.
    """
    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

    def _set_run_font(run_elem):
        rPr = run_elem.find(f"{W}rPr")
        if rPr is None:
            rPr = OxmlElement("w:rPr")
            run_elem.insert(0, rPr)
        rFonts = rPr.find(f"{W}rFonts")
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:cs", "w:eastAsia", "w:hAnsi"):
            rFonts.set(qn(attr), ARIAL_FONT)

    # Iterate all runs in the body
    for run_elem in doc.element.body.iter(qn("w:r")):
        _set_run_font(run_elem)


def apply_zensko_style_to_existing_tables(doc: Document) -> None:
    """Apply Żeńsko style (dark navy header + grey borders) to header-row tables only.

    Skips metadata/key-value tables (Adres obiektu, Firma serwisowa, etc.).
    Heuristic: a table is treated as header-row if its row 0 first cell text
    matches one of HEADER_TABLE_HINTS (case-insensitive).
    """
    for table in doc.tables:
        if not table.rows:
            continue
        first_cell_text = table.rows[0].cells[0].text.strip().lower()
        is_header_row = any(
            first_cell_text.startswith(hint) for hint in HEADER_TABLE_HINTS
        )
        if is_header_row:
            style_table_zensko(table, zebra=True)
        else:
            # Just borders for metadata tables (no header styling).
            _style_borders_only(table)


def _style_borders_only(table) -> None:
    """Apply only thin grey borders (no header shading). For photo grid etc."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:space"), "0")
        b.set(qn("w:color"), TABLE_BORDER_COLOR)
        tblBorders.append(b)
    tblPr.append(tblBorders)


def _make_paragraph_element(text: str, bold: bool = False) -> object:
    """Create a free-standing <w:p> XML element with given text."""
    p_elem = OxmlElement("w:p")
    r_elem = OxmlElement("w:r")
    if bold:
        rPr = OxmlElement("w:rPr")
        b = OxmlElement("w:b")
        rPr.append(b)
        r_elem.append(rPr)
    t_elem = OxmlElement("w:t")
    t_elem.set(qn("xml:space"), "preserve")
    t_elem.text = text
    r_elem.append(t_elem)
    p_elem.append(r_elem)
    return p_elem


def _set_cell_text_bold(cell: _Cell, text: str, bold: bool = True) -> None:
    """Set cell text (single paragraph) and optionally bold."""
    cell.text = text
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = bold


def _populate_data_table(
    table, headers: list[str], rows: list[tuple[str, ...]]
) -> None:
    """Populate a data table: bold header row + data rows."""
    for ci, h in enumerate(headers):
        _set_cell_text_bold(table.cell(0, ci), h, bold=True)
    for ri, row_data in enumerate(rows, start=1):
        for ci, val in enumerate(row_data):
            table.cell(ri, ci).text = val


# ---------- Sekcja VI/IV — Zalecenia (extended) -----------------------------


def update_section_zalecenia(doc: Document) -> None:
    """Replace 2-col 'Zakres czynności / Termin wykonania' table with 6-col Żeńsko-style table.

    Also inserts after it:
      - 'Definicje rodzajów robót remontowych' (heading + 3-col 4-row table)
      - 'Zalecany czas wykonania robót remontowych — stopień pilności' (heading + 3-col 5-row)
      - 'Kryteria oceny i klasyfikacji stanu technicznego' (heading + 3-col 6-row, 5-stopniowa z %)
      - 3 narrative sections (Stan techniczny instalacji ochrony środowiska, Weryfikacja
        kompletności dokumentów, Metody i środki użytkowania)

    Idempotent: detects existing definitions ('Definicje rodzajów robót remontowych') and skips.
    """
    if find_paragraph_by_text(doc, DEFINICJE_ROBOT_HEADER) is not None:
        return  # Already extended

    target_para = find_paragraph_by_text(doc, "Określenie zakresu robót remontowych")
    if target_para is None:
        return

    # Find the 2-col table immediately after target_para
    elem = target_para._element.getnext()
    while elem is not None and elem.tag != qn("w:tbl"):
        elem = elem.getnext()
    if elem is None:
        return
    old_tbl = elem

    # Remove old 2-col table
    old_tbl.getparent().remove(old_tbl)

    # Sequence of inserts after target_para
    current = target_para._element

    # 1. New 6-col zalecenia table (header + 5 empty data rows = 6 rows total)
    zalecenia_table = doc.add_table(rows=6, cols=6)
    for ci, h in enumerate(ZALECENIA_HEADERS):
        _set_cell_text_bold(zalecenia_table.cell(0, ci), h, bold=True)
    # Pre-populate Lp. column with sequential numbers in data rows
    for ri in range(1, 6):
        zalecenia_table.cell(ri, 0).text = str(ri)
    style_table_zensko(zalecenia_table)
    current.addnext(zalecenia_table._element)
    current = zalecenia_table._element

    # 2. Definicje rodzajów robót — heading + table (3 cols × 4 rows: 1 header + 3 data)
    def_header = _make_paragraph_element(DEFINICJE_ROBOT_HEADER, bold=True)
    current.addnext(def_header)
    current = def_header

    rodzaje_table = doc.add_table(rows=4, cols=3)
    _populate_data_table(rodzaje_table, RODZAJE_ROBOT_HEADERS, RODZAJE_ROBOT_ROWS)
    style_table_zensko(rodzaje_table)
    current.addnext(rodzaje_table._element)
    current = rodzaje_table._element

    # 3. Stopnie pilności — heading + table (3 cols × 5 rows: 1 header + 4 data)
    pil_header = _make_paragraph_element(STOPNIE_PILNOSCI_HEADER, bold=True)
    current.addnext(pil_header)
    current = pil_header

    pil_table = doc.add_table(rows=5, cols=3)
    _populate_data_table(pil_table, STOPNIE_PILNOSCI_HEADERS, STOPNIE_PILNOSCI_ROWS)
    style_table_zensko(pil_table)
    current.addnext(pil_table._element)
    current = pil_table._element

    # 4. Kryteria klasyfikacji (5-stopniowa z %) — heading + table (3 cols × 6 rows)
    kry_header = _make_paragraph_element(KRYTERIA_KLASYFIKACJI_HEADER, bold=True)
    current.addnext(kry_header)
    current = kry_header

    kry_table = doc.add_table(rows=6, cols=3)
    _populate_data_table(kry_table, KRYTERIA_KLASYFIKACJI_HEADERS, KRYTERIA_KLASYFIKACJI_ROWS)
    style_table_zensko(kry_table)
    current.addnext(kry_table._element)
    current = kry_table._element

    # 5. Narrative sections (3) — heading + placeholder paragraph each
    for nar_header, nar_placeholder in NARRATYWNE_SECTIONS:
        hp = _make_paragraph_element(nar_header, bold=True)
        current.addnext(hp)
        current = hp
        pp = _make_paragraph_element(nar_placeholder, bold=False)
        current.addnext(pp)
        current = pp


# ---------- Sekcja VII — siatka fotografii ----------------------------------


def update_section_vii_photos(doc: Document) -> None:
    """Replace dotted-line placeholders in section VII with a structured empty photo grid.

    Inserts subtitle before the grid:
      'Dokumentacja fotograficzna wykonana podczas kontroli (elementy ... do remontu)'

    Idempotent — removes any previously-inserted subtitle + photo grid before re-adding.
    """
    # Remove existing subtitle paragraph(s) if present (from previous runs)
    existing_subtitle = find_paragraph_by_text(
        doc, "Dokumentacja fotograficzna wykonana podczas kontroli"
    )
    if existing_subtitle is not None:
        # Also remove the photo table immediately following it (if any)
        next_elem = existing_subtitle._element.getnext()
        if next_elem is not None and next_elem.tag == qn("w:tbl"):
            next_elem.getparent().remove(next_elem)
        existing_subtitle._element.getparent().remove(existing_subtitle._element)

    numeracja_para = find_paragraph_by_text(doc, "Numerację fotografii")
    if numeracja_para is None:
        return

    # Determine which paragraph ends section VII (next section header)
    # In roczny: VII. PODPISY
    # In 5-letni: VIII. PODPISY
    end_marker = None
    for marker in ("VIII. PODPISY", "VII. PODPISY"):
        if find_paragraph_by_text(doc, marker) is not None:
            end_marker = marker
            break
    if end_marker is None:
        return

    # Remove all dotted-line / empty paragraphs between numeracja_para and end_marker
    elem = numeracja_para._element.getnext()
    while elem is not None:
        if elem.tag == qn("w:p"):
            text = "".join(t.text or "" for t in elem.iter(qn("w:t")))
            if end_marker in text:
                break  # Reached the end marker — stop
            stripped = text.strip()
            # Match dotted placeholder (many . or … chars) or empty paragraph
            dot_count = sum(1 for c in stripped if c in ".…")
            is_dotted = dot_count >= 10 and dot_count >= len(stripped) * 0.8
            if not stripped or is_dotted:
                next_elem = elem.getnext()
                elem.getparent().remove(elem)
                elem = next_elem
                continue
        elif elem.tag == qn("w:tbl"):
            break  # Don't go past a table
        elem = elem.getnext()

    # Now insert subtitle + photo grid after numeracja_para
    current = numeracja_para._element

    subtitle = _make_paragraph_element(SEKCJA_VII_SUBTITLE, bold=True)
    current.addnext(subtitle)
    current = subtitle

    photo_table = doc.add_table(rows=PHOTO_GRID_ROWS, cols=PHOTO_GRID_COLS)
    # Empty cells — inspector wkleja zdjęcia bezpośrednio do każdej komórki.
    # Photo grid — just borders, no header row styling.
    _style_borders_only(photo_table)
    current.addnext(photo_table._element)


def replace_text_in_cell(cell: _Cell, old: str, new: str) -> None:
    """Replace text in a cell, preserving run-level formatting where possible."""
    for p in cell.paragraphs:
        if old not in p.text:
            continue
        for run in p.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                return
        # Fallback: text spans multiple runs
        full = p.text
        new_full = full.replace(old, new)
        for r in p.runs[1:]:
            r._element.getparent().remove(r._element)
        if p.runs:
            p.runs[0].text = new_full
        return


def generate_roczny_R() -> None:
    """Roczny rozszerzona — modify E3, E6, E7 + intro + serwis note (style-preserving)."""
    doc = Document(str(ROCZNY_SRC))
    table = doc.tables[6]  # Table #6 = III. USTALENIA

    # Templates from Element 1 fundament C1 (row 2, col 1) — unchanged, preserves original styling.
    templates = extract_templates(table.cell(2, 1))

    # Element 3 wieża → row 6
    replace_cell_styled(table.cell(6, 1), ROCZNY_R_E3_C1, templates)
    # Element 6 wirnik → row 12
    replace_cell_styled(table.cell(12, 1), ROCZNY_R_E6_C1, templates)
    # Element 7 piasta → row 14
    replace_cell_styled(table.cell(14, 1), ROCZNY_R_E7_C1, templates)

    add_intro_above_table_iii(doc, INTRO_ROCZNY_R)
    remove_serwis_checklist(doc)
    add_serwis_note(doc)
    update_section_zalecenia(doc)
    update_section_vii_photos(doc)

    # Global Żeńsko styling — font + table styling (header tables only)
    fix_section_header_rows(doc)
    apply_arial_font_globally(doc)

    doc.save(str(ROCZNY_R_OUT))
    print(f"✓ {ROCZNY_R_OUT.name}")


def generate_roczny_U() -> None:
    """Roczny uproszczona — modify all 15 element C1 cells + intro (style-preserving)."""
    doc = Document(str(ROCZNY_SRC))
    table = doc.tables[6]

    # Templates from Element 1 fundament C1 — unchanged in source ROCZNY_R.
    templates = extract_templates(table.cell(2, 1))

    for elem_k in range(1, 16):
        data_row = 2 * elem_k
        replace_cell_styled(table.cell(data_row, 1), ROCZNY_U[elem_k], templates)

    # Uproszczona — bez ⚙, więc bez legendy i bez serwis note (info o serwisie jest w intro).
    remove_paragraphs_with_markers(doc, (SERWIS_MARKER,))
    add_intro_above_table_iii(doc, INTRO_ROCZNY_U, include_legenda=False)
    remove_serwis_checklist(doc)
    update_section_zalecenia(doc)
    update_section_vii_photos(doc)

    fix_section_header_rows(doc)
    apply_arial_font_globally(doc)

    doc.save(str(ROCZNY_U_OUT))
    print(f"✓ {ROCZNY_U_OUT.name}")


def generate_5letni() -> None:
    """5-letni — modify all 16 element C2 cells + 3 element C1 cells (E3/E6/E7) + intro
    + serwis note (style-preserving)."""
    doc = Document(str(PIECIOLETNI_SRC))
    table = doc.tables[8]  # Table #8 = III. USTALENIA

    # Templates from Element 16 estetyka C2 (row 32, col 2) — unchanged, preserves original styling.
    templates_c2 = extract_templates(table.cell(32, 2))
    # Templates for C1 — use Element 1 fundament C1 (unchanged)
    templates_c1 = extract_templates(table.cell(2, 1))

    # C2 — all 16 elements
    for elem_k in range(1, 17):
        data_row = 2 * elem_k
        replace_cell_styled(table.cell(data_row, 2), P5_C2[elem_k], templates_c2)

    # C1 (Zakres roczny poszerzony) — E3 wieża, E6 wirnik, E7 piasta (usuwa „dron",
    # warunkuje analizy SCADA, oznaczę ⚙ czynności serwisowe).
    replace_cell_styled(table.cell(6, 1), P5_E3_C1, templates_c1)
    replace_cell_styled(table.cell(12, 1), P5_E6_C1, templates_c1)
    replace_cell_styled(table.cell(14, 1), P5_E7_C1, templates_c1)

    add_intro_above_table_iii(doc, INTRO_5LETNI)
    remove_serwis_checklist(doc)
    add_serwis_note(doc)
    apply_pomiar_template(doc)
    update_section_zalecenia(doc)
    update_section_vii_photos(doc)

    fix_section_header_rows(doc)
    apply_arial_font_globally(doc)

    doc.save(str(PIECIOLETNI_OUT))
    print(f"✓ {PIECIOLETNI_OUT.name}")


def generate_5letni_O() -> None:
    """5-letni odrębny — wariant dla osobnych protokołów (rocznego + 5-letniego).

    Różnice względem domyślnego 5-letniego:
      - Brak nakładki UWAGA o pełnym zakresie kontroli rocznej (usuwana tabela 1x1 na początku)
      - Kolumna C1 'ZAKRES KONTROLI ROCZNEJ (poszerzony)' usunięta z tabeli III
      - Nagłówek kolumny C2 zmieniony z 'ZAKRES DODATKOWY 5-LETNI' na 'ZAKRES 5-LETNI'
      - Nota wstępna (intro) wyjaśnia wariant odrębny
    """
    doc = Document(str(PIECIOLETNI_OUT))  # Source = updated combined 5-letni
    table = doc.tables[8]

    templates = extract_templates(table.cell(32, 2))

    # Re-apply C2 modifications (idempotent — same content as combined)
    for elem_k in range(1, 17):
        data_row = 2 * elem_k
        replace_cell_styled(table.cell(data_row, 2), P5_C2[elem_k], templates)

    # Replace intro with the odrębny variant
    add_intro_above_table_iii(doc, INTRO_5LETNI_O)
    remove_serwis_checklist(doc)
    add_serwis_note(doc)
    apply_pomiar_template(doc)
    update_section_zalecenia(doc)
    update_section_vii_photos(doc)

    # Update the C2 header cell text BEFORE removing column C1.
    replace_text_in_cell(table.cell(0, 2), "ZAKRES DODATKOWY 5-LETNI", "ZAKRES 5-LETNI")

    # Remove column C1 (Zakres roczny poszerzony) from table III.
    remove_column_from_table(table, 1)

    # Remove the UWAGA table about full annual scope inclusion.
    remove_uwaga_combined_scope(doc)

    # Shorten the description of tabela III columns (no more 'Zakres roczny').
    update_table_iii_description_for_O(doc)

    fix_section_header_rows(doc)
    apply_arial_font_globally(doc)

    doc.save(str(PIECIOLETNI_O_OUT))
    print(f"✓ {PIECIOLETNI_O_OUT.name}")


def update_raport_zmian() -> None:
    """Append section 9 to Raport_zmian_wzory_PIIB.docx (idempotent — removes existing section 9 first)."""
    doc = Document(str(RAPORT_SRC))

    # Idempotency: remove any existing paragraphs that belong to a previously-added section 9.
    body = doc.element.body
    paragraphs = list(doc.paragraphs)
    section_9_marker = "9. Aktualizacja zakresów kontroli z dnia 2026-05-10"
    first_section_9_idx = next(
        (i for i, p in enumerate(paragraphs) if p.text.strip().startswith(section_9_marker)),
        None,
    )
    if first_section_9_idx is not None:
        for p in paragraphs[first_section_9_idx:]:
            p._element.getparent().remove(p._element)

    body_paragraphs = [
        ("9. Aktualizacja zakresów kontroli z dnia 2026-05-10", True),
        (
            "W ramach dostosowania zakresów kontroli do realiów rynku polskich kontroli okresowych turbin "
            "wiatrowych (kontrola roczna w przedziale 300–500 zł netto, kontrola pięcioletnia w przedziale "
            "1000–1500 zł netto z pomiarami elektrycznymi) wprowadzono następujące zmiany w sekcji III. "
            "USTALENIA obu wzorów:",
            False,
        ),
        ("9.1. Wprowadzono dwa warianty protokołu rocznego", True),
        (
            "Inspektor wybiera wariant zgodnie z wcześniejszymi ustaleniami z klientem. Pliki w katalogu "
            "wzory_PIIB/:",
            False,
        ),
        (
            "•  Protokol_Kontroli_Rocznej_EW_PIIB_R.docx — wariant rozszerzony (z wjazdem na konstrukcję). "
            "Zawiera oględziny zewnętrzne, wewnętrzne wieży, gondoli (po wjeździe), wirnika z gondoli, "
            "łożysk, instalacji.",
            False,
        ),
        (
            "•  Protokol_Kontroli_Rocznej_EW_PIIB_U.docx — wariant uproszczony (kontrola bez wjazdu). "
            "Zawiera oględziny z poziomu terenu (lornetka), inspekcję wewnętrzną pierwszego segmentu "
            "wieży, weryfikację dokumentacji i analizę SCADA. Czynności wymagające wjazdu są w gestii "
            "serwisu producenta lub osobnego zlecenia.",
            False,
        ),
        ("9.2. Protokół 5-letni — dwa warianty, zawsze z wjazdem", True),
        (
            "Ze względu na charakter inspekcji 5-letniej (przydatność do użytkowania, art. 62 ust. 1 pkt 2 "
            "PB) oraz konieczność wykonania pomiarów elektrycznych i odgromowych — kontrola 5-letnia "
            "zawsze obejmuje wjazd na konstrukcję. Pliki:",
            False,
        ),
        (
            "•  Protokol_Kontroli_5-letniej_EW_PIIB.docx — wariant domyślny (połączony z kontrolą roczną). "
            "Zawiera kolumnę „ZAKRES KONTROLI ROCZNEJ (poszerzony)” oraz „ZAKRES DODATKOWY 5-LETNI” "
            "i obejmuje pełen zakres rocznej + dodatkowy 5-letni w jednym dokumencie.",
            False,
        ),
        (
            "•  Protokol_Kontroli_5-letniej_EW_PIIB_O.docx — wariant odrębny (osobny protokół 5-letni "
            "wymagany przez niektóre PINB-y). Bez kolumny „ZAKRES KONTROLI ROCZNEJ” (kontrola roczna "
            "sporządzana w odrębnym protokole), nagłówek kolumny zmieniony na „ZAKRES 5-LETNI”, usunięta "
            "uwaga o objęciu pełnego zakresu rocznej.",
            False,
        ),
        ("9.3. Wprowadzono podział czynności w tabeli III", True),
        (
            "W kolumnach OPIS STANU TECHNICZNEGO (roczny) i ZAKRES DODATKOWY 5-LETNI (5-letni) "
            "rozdzielono czynności wykonywane w ramach inspekcji okresowej (📋) od czynności "
            "specjalistycznych realizowanych przez serwis producenta lub w ramach osobnego zlecenia (⚙). "
            "Dodatkowo wprowadzono piktogram ℹ dla czynności realizowanych przez podmioty zewnętrzne "
            "(np. badania techniczne UDT).",
            False,
        ),
        ("9.4. Kluczowe wycięcia z zakresu obowiązkowego", True),
        (
            "•  Drony / dostęp linowy do łopat — przeniesione do ⚙ (wcześniej w zakresie rocznym i 5-letnim).",
            False,
        ),
        (
            "•  Pomiar luzu łożyska wieńcowego, diagnostyka wibracyjna łożyska głównego, kontrola "
            "momentów dokręcenia 100% — przeniesione do ⚙ (serwis producenta wg art. 8b u.w.).",
            False,
        ),
        (
            "•  Badania NDT spawów flansz / spoin ramy nośnej, pomiar grubości powłok grubościomierzem, "
            "badania nieniszczące betonu, próby obciążeniowe podestów / kotwiczeń schodów — przeniesione "
            "do ⚙.",
            False,
        ),
        (
            "•  Termowizja rozdzielnic w protokole 5-letnim — pozostaje w zakresie inspekcji, ale jako "
            "czynność OPCJONALNA (checkbox ☐) zaznaczana przez inspektora po wykonaniu.",
            False,
        ),
        (
            "•  Analiza danych SCADA (alarmy, temperatury, drgania) opatrzona zastrzeżeniem „jeżeli "
            "udostępniono dostęp do danych” — wcześniej zapis bezwarunkowy.",
            False,
        ),
        ("9.5. Pozycje pozostawione w zakresie inspekcji okresowej", True),
        (
            "•  Pomiary elektryczne 5-letnie (rezystancja izolacji, pętla zwarcia, ciągłość PE, RCD) — "
            "obowiązkowo, zgodnie z art. 62 ust. 1 pkt 2 PB.",
            False,
        ),
        (
            "•  Pomiary instalacji odgromowej (rezystancja uziemienia, ciągłość LPS, sprawność SPD) — "
            "obowiązkowo.",
            False,
        ),
        (
            "•  Ocena estetyki w protokole 5-letnim — bez zmian.",
            False,
        ),
        (
            "•  Weryfikacja dokumentów UDT, dat ważności gaśnic, planu ewakuacji — bez zmian.",
            False,
        ),
        ("9.6. Uzupełnienie sekcji V. Informacje o serwisie technicznym turbiny", True),
        (
            "Pod akapitem opisowym dodano notę informującą wprost, że czynności specjalistyczne oznaczone "
            "piktogramem ⚙ w tabeli III są realizowane przez serwis producenta zgodnie z umową serwisową.",
            False,
        ),
        ("10. Aktualizacja sekcji VI/IV (Zalecenia) i VII/VI (Dokumentacja) z dnia 2026-05-11", True),
        (
            "Dostosowanie struktury sekcji VI (5-letni) / IV (roczny) ZALECENIA oraz sekcji VII/VI "
            "DOKUMENTACJA do wzoru aktualnie generowanego przez aplikację (oraz protokołu referencyjnego "
            "Żeńsko 003/P/2026):",
            False,
        ),
        ("10.1. Tabela „Zakres robót remontowych i kolejność wykonywania” — rozszerzenie", True),
        (
            "Zastąpiono dotychczasową tabelę 2-kolumnową („Zakres czynności / Termin wykonania”) tabelą "
            "6-kolumnową: Lp. / Element / lokalizacja / Zakres robót remontowych / Rodzaj / Pilność / "
            "Termin wykonania. Domyślnie 5 wierszy danych z prepopulowanymi numerami Lp. (1–5).",
            False,
        ),
        ("10.2. Dodano tabelę „Definicje rodzajów robót remontowych”", True),
        (
            "Klasyfikacja K (Konserwacja) / NB (Naprawa bieżąca) / NG (Naprawa główna) — 3-kolumnowa "
            "tabela definicji wstawiona bezpośrednio pod tabelą zaleceń. Spójna z notacją używaną przez "
            "aplikację oraz wcześniejsze protokoły Prowatech.",
            False,
        ),
        ("10.3. Dodano tabelę „Stopnie pilności” (Zalecany czas wykonania)", True),
        (
            "4 stopnie pilności: I (natychmiast) / II (do 3 miesięcy) / III (do 12 miesięcy) / IV (do 5 "
            "lat) — 3-kolumnowa tabela z opisami zaczerpnięta ze wzoru aplikacyjnego.",
            False,
        ),
        ("10.4. Dodano tabelę „Kryteria oceny i klasyfikacji stanu technicznego” (5-stopniowa z %)", True),
        (
            "Klasyfikacja stanu technicznego z zużyciem procentowym: dobry (0–15%) / zadowalający "
            "(16–30%) / średni (31–50%) / zły (51–70%) / awaryjny (>71%). Tabela ta jest LEGENDĄ "
            "klasyfikacyjną — uzupełnia 4-stopniową skalę OCENA (dobry/dostateczny/niedostateczny/"
            "awaryjny) używaną w sekcji III (tabela ustaleń) oraz w „Przyjętych kryteriach oceny” na "
            "początku protokołu. Obie skale współistnieją zgodnie ze wzorem aplikacyjnym i Żeńsko.",
            False,
        ),
        ("10.5. Dodano 3 sekcje narracyjne", True),
        (
            "Pod tabelami definicji wstawiono trzy sekcje opisowe (zaczerpnięte z Żeńsko):",
            False,
        ),
        (
            "•  Stan techniczny instalacji ochrony środowiska",
            False,
        ),
        (
            "•  Weryfikacja kompletności i aktualności dokumentów",
            False,
        ),
        (
            "•  Metody i środki użytkowania elementów narażonych na szkodliwe wpływy atmosferyczne i "
            "niszczące działanie innych czynników",
            False,
        ),
        ("10.6. Sekcja VII/VI DOKUMENTACJA — siatka fotografii", True),
        (
            "Zastąpiono linie kropek strukturowaną tabelą 6×2 (12 placeholderów „Zdjęcie nr N — wstaw "
            "fotografię”). Dodano podtytuł „Dokumentacja fotograficzna wykonana podczas kontroli "
            "(elementy obiektu posiadające usterki lub wady, przewidziane do remontu)”.",
            False,
        ),
        ("10.7. Zakres zmian", True),
        (
            "Powyższe zmiany zastosowane do wszystkich czterech wzorów: Protokol_Kontroli_Rocznej_EW_PIIB_R, "
            "Protokol_Kontroli_Rocznej_EW_PIIB_U, Protokol_Kontroli_5-letniej_EW_PIIB (połączony) oraz "
            "Protokol_Kontroli_5-letniej_EW_PIIB_O (odrębny).",
            False,
        ),
        ("10.8. Skala OCENA w sekcji III — bez zmian", True),
        (
            "Skala 4-stopniowa PIIB (dobry / dostateczny / niedostateczny / awaryjny) w kolumnie OCENA "
            "tabeli III pozostała bez zmian — zgodnie z wzorem aplikacyjnym oraz Żeńsko. Nowa 5-stopniowa "
            "skala z procentami zużycia w sekcji VI jest LEGENDĄ KLASYFIKACYJNĄ uzupełniającą, nie "
            "zastępującą skalę PIIB w ustaleniach.",
            False,
        ),
        ("11. Poprawa czytelności nagłówków elementów w tabeli III z dnia 2026-05-11", True),
        (
            "W tabeli III. USTALENIA wiersze rozdzielające poszczególne elementy konstrukcyjne "
            "(np. „1. FUNDAMENT I POSADOWIENIE”, „6. WIRNIK / ROTOR Z ŁOPATAMI”) miały biały bold "
            "tekst na jasnoszarym tle zebry (HEX F5F7F9), przez co były praktycznie nieczytelne. "
            "Zmieniono kolor wypełnienia tych wierszy na granatowy (HEX 1B2230 — graphite800, "
            "spójnie z nagłówkami kolumn tabeli oraz tokenami stylów aplikacji), zachowując biały "
            "bold tekst. Zmiana wyłącznie wizualna — bez modyfikacji treści ani struktury tabeli.",
            False,
        ),
        (
            "Zakres: wszystkie cztery wzory (Protokol_Kontroli_Rocznej_EW_PIIB_R, "
            "Protokol_Kontroli_Rocznej_EW_PIIB_U, Protokol_Kontroli_5-letniej_EW_PIIB, "
            "Protokol_Kontroli_5-letniej_EW_PIIB_O).",
            False,
        ),
    ]

    for text, bold in body_paragraphs:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = bold

    doc.save(str(RAPORT_SRC))
    print(f"✓ {RAPORT_SRC.name} (zaktualizowany)")


def main() -> None:
    generate_roczny_R()
    generate_roczny_U()
    generate_5letni()
    generate_5letni_O()
    update_raport_zmian()
    print("\nDone.")


if __name__ == "__main__":
    main()
